import os
import tempfile
import zipfile
from typing import List, Optional, Tuple
import fitz  # PyMuPDF
from PIL import Image
import pdf2image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
# import pdfplumber  # Removed due to compatibility issues
import re
import io

class DocumentConverter:
    """Handles document format conversions"""
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def pdf_to_word(self, pdf_path: str, force_english_numbers: bool = True, preserve_formatting: bool = True, ocr_language: str = "eng") -> Optional[str]:
        """Convert PDF to Word document"""
        try:
            # Create new Word document
            doc = Document()
            
            # Extract text from PDF using PyMuPDF
            pdf_doc = fitz.open(pdf_path)
            for page_num in range(len(pdf_doc)):
                page = pdf_doc.load_page(page_num)
                
                # Add page break for pages after the first
                if page_num > 0:
                    doc.add_page_break()
                
                # Add page header
                doc.add_heading(f'Page {page_num + 1}', level=2)
                
                # Extract text
                page_text = page.get_text()
                
                if page_text:
                    # Force English numbers if requested
                    if force_english_numbers:
                        page_text = self.force_english_numbers(page_text)
                    
                    # Split text into paragraphs and add to document
                    paragraphs = page_text.split('\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            p = doc.add_paragraph(paragraph)
                            if preserve_formatting:
                                # Try to preserve some basic formatting
                                if paragraph.isupper():
                                    p.runs[0].bold = True
                                if len(paragraph) < 50 and paragraph.count(' ') < 5:
                                    p.runs[0].bold = True
            
            pdf_doc.close()
            
            # Save Word document
            output_path = os.path.join(self.temp_dir, f"converted_word_{int(time.time())}.docx")
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error converting PDF to Word: {str(e)}")
    
    def word_to_pdf(self, word_path: str) -> Optional[str]:
        """Convert Word document to PDF"""
        try:
            # Try to use docx2pdf if available, otherwise use python-docx with reportlab
            try:
                from docx2pdf import convert
                output_path = os.path.join(self.temp_dir, f"converted_pdf_{int(time.time())}.pdf")
                convert(word_path, output_path)
                return output_path
            except ImportError:
                # Fallback method using reportlab
                return self._word_to_pdf_fallback(word_path)
                
        except Exception as e:
            raise Exception(f"Error converting Word to PDF: {str(e)}")
    
    def _word_to_pdf_fallback(self, word_path: str) -> Optional[str]:
        """Fallback method for Word to PDF conversion using reportlab"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            
            # Read Word document
            doc = Document(word_path)
            
            # Create PDF
            output_path = os.path.join(self.temp_dir, f"converted_pdf_{int(time.time())}.pdf")
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            y_position = height - inch
            line_height = 14
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Check if we need a new page
                    if y_position < inch:
                        c.showPage()
                        y_position = height - inch
                    
                    # Add text to PDF
                    text = paragraph.text
                    
                    # Simple text wrapping
                    max_width = width - 2 * inch
                    words = text.split()
                    lines = []
                    current_line = []
                    
                    for word in words:
                        test_line = ' '.join(current_line + [word])
                        if c.stringWidth(test_line) < max_width:
                            current_line.append(word)
                        else:
                            if current_line:
                                lines.append(' '.join(current_line))
                                current_line = [word]
                            else:
                                lines.append(word)
                    
                    if current_line:
                        lines.append(' '.join(current_line))
                    
                    # Add lines to PDF
                    for line in lines:
                        c.drawString(inch, y_position, line)
                        y_position -= line_height
            
            c.save()
            return output_path
            
        except Exception as e:
            raise Exception(f"Error in fallback Word to PDF conversion: {str(e)}")
    
    def pdf_to_images(self, pdf_path: str, dpi: int = 300, format: str = "png", page_range: Optional[str] = None) -> List[str]:
        """Convert PDF pages to images"""
        try:
            # Parse page range if provided
            first_page = None
            last_page = None
            
            if page_range:
                if '-' in page_range:
                    start, end = page_range.split('-')
                    first_page = int(start.strip())
                    last_page = int(end.strip())
                else:
                    first_page = int(page_range.strip())
                    last_page = first_page
            
            # Convert PDF to images
            images = pdf2image.convert_from_path(
                pdf_path,
                dpi=dpi,
                output_format=format.upper(),
                first_page=first_page,
                last_page=last_page,
                thread_count=2
            )
            
            # Save images
            image_paths = []
            for i, image in enumerate(images):
                page_num = i + 1 if not first_page else first_page + i
                output_path = os.path.join(
                    self.temp_dir, 
                    f"page_{page_num:03d}_{int(time.time())}.{format}"
                )
                
                if format.lower() == "jpeg" and image.mode == "RGBA":
                    # Convert RGBA to RGB for JPEG
                    rgb_image = Image.new("RGB", image.size, (255, 255, 255))
                    rgb_image.paste(image, mask=image.split()[-1])
                    rgb_image.save(output_path, format.upper(), quality=95, optimize=True)
                else:
                    image.save(output_path, format.upper(), optimize=True)
                
                image_paths.append(output_path)
            
            return image_paths
            
        except Exception as e:
            raise Exception(f"Error converting PDF to images: {str(e)}")
    
    def images_to_pdf(self, image_paths: List[str], page_size: str = "A4", orientation: str = "portrait") -> Optional[str]:
        """Convert images to PDF"""
        try:
            # Define page sizes in points (1 inch = 72 points)
            page_sizes = {
                "A4": (595, 842),
                "Letter": (612, 792),
                "A3": (842, 1191),
                "Custom": None
            }
            
            if orientation == "landscape":
                width, height = page_sizes[page_size]
                page_sizes[page_size] = (height, width)
            
            # Create PDF document
            doc = fitz.open()
            
            for image_path in image_paths:
                # Open image
                img = Image.open(image_path)
                
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Get page dimensions
                if page_size == "Custom":
                    # Use image size
                    page_width = img.width * 72 / 96  # Assuming 96 DPI
                    page_height = img.height * 72 / 96
                else:
                    page_width, page_height = page_sizes[page_size]
                
                # Create new page
                page = doc.new_page(width=page_width, height=page_height)
                
                # Calculate image placement to fit page while maintaining aspect ratio
                img_width, img_height = img.size
                img_aspect = img_width / img_height
                page_aspect = page_width / page_height
                
                if img_aspect > page_aspect:
                    # Image is wider relative to page
                    scaled_width = page_width * 0.95  # Leave some margin
                    scaled_height = scaled_width / img_aspect
                else:
                    # Image is taller relative to page
                    scaled_height = page_height * 0.95  # Leave some margin
                    scaled_width = scaled_height * img_aspect
                
                # Center the image
                x = (page_width - scaled_width) / 2
                y = (page_height - scaled_height) / 2
                
                # Insert image
                img_rect = fitz.Rect(x, y, x + scaled_width, y + scaled_height)
                
                # Save image temporarily in a format PyMuPDF can handle
                temp_img_path = os.path.join(self.temp_dir, f"temp_img_{int(time.time())}.png")
                img.save(temp_img_path, "PNG")
                
                page.insert_image(img_rect, filename=temp_img_path)
                
                # Clean up temporary image
                os.remove(temp_img_path)
            
            # Save PDF
            output_path = os.path.join(self.temp_dir, f"images_to_pdf_{int(time.time())}.pdf")
            doc.save(output_path)
            doc.close()
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error converting images to PDF: {str(e)}")
    
    def create_images_zip(self, image_paths: List[str], original_filename: str) -> str:
        """Create a ZIP file containing all images"""
        try:
            zip_path = os.path.join(self.temp_dir, f"{original_filename}_images_{int(time.time())}.zip")
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for i, image_path in enumerate(image_paths):
                    # Get file extension
                    ext = os.path.splitext(image_path)[1]
                    filename = f"page_{i+1:03d}{ext}"
                    zipf.write(image_path, filename)
            
            return zip_path
            
        except Exception as e:
            raise Exception(f"Error creating images ZIP: {str(e)}")
    
    def force_english_numbers(self, text: str) -> str:
        """Convert Hindi/Devanagari numbers to English digits"""
        # Devanagari to English number mapping
        hindi_to_english = {
            '०': '0', '१': '1', '२': '2', '३': '3', '४': '4',
            '५': '5', '६': '6', '७': '7', '८': '8', '९': '9'
        }
        
        # Arabic-Indic to English number mapping
        arabic_indic_to_english = {
            '٠': '0', '١': '1', '٢': '2', '٣': '3', '٤': '4',
            '٥': '5', '٦': '6', '٧': '7', '٨': '8', '٩': '9'
        }
        
        # Apply conversions
        for hindi_digit, english_digit in hindi_to_english.items():
            text = text.replace(hindi_digit, english_digit)
        
        for arabic_digit, english_digit in arabic_indic_to_english.items():
            text = text.replace(arabic_digit, english_digit)
        
        return text

import time
