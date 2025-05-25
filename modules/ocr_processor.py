import os
import tempfile
import re
from typing import Optional, List
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import pdf2image
from docx import Document
import io

class OCRProcessor:
    """Handles OCR processing for PDF documents"""
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
        
        # Configure tesseract path if needed (common Windows path)
        # If tesseract is not in PATH, uncomment and modify this line:
        # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    
    def detect_language(self, image_path: str) -> str:
        """Auto-detect language in image using OCR"""
        try:
            # Get language detection from tesseract
            osd = pytesseract.image_to_osd(image_path, output_type=pytesseract.Output.DICT)
            
            # For simplicity, we'll detect based on character patterns
            text_sample = pytesseract.image_to_string(image_path, lang='eng+hin')[:500]
            
            # Check for Hindi/Devanagari characters
            hindi_pattern = re.compile(r'[\u0900-\u097F]')
            english_pattern = re.compile(r'[a-zA-Z]')
            
            hindi_matches = len(hindi_pattern.findall(text_sample))
            english_matches = len(english_pattern.findall(text_sample))
            
            if hindi_matches > english_matches * 0.3:  # Significant Hindi content
                return "hin+eng"
            else:
                return "eng"
                
        except Exception as e:
            # Default to English if detection fails
            return "eng"
    
    def get_tesseract_language(self, language: str) -> str:
        """Convert language parameter to tesseract language codes"""
        language_map = {
            "auto_detect": "eng+hin",
            "english": "eng",
            "hindi": "hin",
            "english_+_hindi": "eng+hin"
        }
        return language_map.get(language, "eng")
    
    def force_english_numbers(self, text: str) -> str:
        """Convert Hindi/Devanagari numbers to English digits"""
        # Devanagari to English number mapping
        hindi_to_english = {
            '०': '0', '१': '1', '२': '2', '३': '3', '४': '4',
            '५': '5', '६': '6', '७': '7', '८': '8', '९': '9'
        }
        
        # Arabic-Indic to English number mapping (used in some Hindi contexts)
        arabic_indic_to_english = {
            '٠': '0', '١': '1', '٢': '2', '٣': '3', '٤': '4',
            '٥': '5', '٦': '6', '٧': '7', '٨': '8', '٩': '9'
        }
        
        # Apply conversions
        for hindi_digit, english_digit in hindi_to_english.items():
            text = text.replace(hindi_digit, english_digit)
        
        for arabic_digit, english_digit in arabic_indic_to_english.items():
            text = text.replace(arabic_digit, english_digit)
        
        return text
    
    def process_pdf_ocr(self, pdf_path: str, language: str = "auto_detect", force_english_numbers: bool = True) -> Optional[str]:
        """Process PDF with OCR to make it searchable"""
        try:
            # Convert PDF pages to images
            images = pdf2image.convert_from_path(
                pdf_path,
                dpi=300
            )
            
            # Create new PDF with OCR text
            doc = fitz.open()
            
            for i, image in enumerate(images):
                # Save image temporarily
                img_path = os.path.join(self.temp_dir, f"temp_page_{i}_{int(time.time())}.png")
                image.save(img_path)
                
                try:
                    # Auto-detect language if needed
                    if language == "auto_detect":
                        detected_lang = self.detect_language(img_path)
                        tesseract_lang = detected_lang
                    else:
                        tesseract_lang = self.get_tesseract_language(language)
                    
                    # Perform OCR with better configuration
                    custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz।॥ःं०१२३४५६७८९ '
                    
                    # Get OCR text and bounding boxes
                    ocr_data = pytesseract.image_to_data(
                        img_path,
                        lang=tesseract_lang,
                        config=custom_config,
                        output_type=pytesseract.Output.DICT
                    )
                    
                    # Create new page
                    page_width, page_height = image.size
                    # Convert pixels to points (PDF uses points: 1 inch = 72 points)
                    page_width_pt = page_width * 72 / 300  # 300 DPI
                    page_height_pt = page_height * 72 / 300
                    
                    new_page = doc.new_page(width=page_width_pt, height=page_height_pt)
                    
                    # Insert the original image as background
                    img_rect = fitz.Rect(0, 0, page_width_pt, page_height_pt)
                    new_page.insert_image(img_rect, filename=img_path)
                    
                    # Add invisible OCR text layer
                    for j in range(len(ocr_data['text'])):
                        if int(ocr_data['conf'][j]) > 30:  # Confidence threshold
                            text = ocr_data['text'][j].strip()
                            if text:
                                # Force English numbers if requested
                                if force_english_numbers:
                                    text = self.force_english_numbers(text)
                                
                                # Convert pixel coordinates to PDF points
                                x = ocr_data['left'][j] * 72 / 300
                                y = page_height_pt - (ocr_data['top'][j] * 72 / 300)  # PDF coordinates are bottom-up
                                width = ocr_data['width'][j] * 72 / 300
                                height = ocr_data['height'][j] * 72 / 300
                                
                                # Calculate font size based on height
                                font_size = max(height * 0.8, 1)
                                
                                # Insert invisible text
                                text_rect = fitz.Rect(x, y - height, x + width, y)
                                new_page.insert_text(
                                    (x, y),
                                    text,
                                    fontsize=font_size,
                                    color=(1, 1, 1),  # White (invisible)
                                    overlay=False
                                )
                
                except Exception as e:
                    print(f"Error processing page {i}: {str(e)}")
                    # Add page without OCR as fallback
                    page_width, page_height = image.size
                    page_width_pt = page_width * 72 / 300
                    page_height_pt = page_height * 72 / 300
                    new_page = doc.new_page(width=page_width_pt, height=page_height_pt)
                    img_rect = fitz.Rect(0, 0, page_width_pt, page_height_pt)
                    new_page.insert_image(img_rect, filename=img_path)
                
                finally:
                    # Clean up temporary image
                    if os.path.exists(img_path):
                        os.remove(img_path)
            
            # Save the searchable PDF
            output_path = os.path.join(self.temp_dir, f"searchable_pdf_{int(time.time())}.pdf")
            doc.save(output_path)
            doc.close()
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error processing PDF with OCR: {str(e)}")
    
    def extract_text_with_ocr(self, pdf_path: str, language: str = "auto_detect", output_format: str = "plain text") -> Optional[str]:
        """Extract text from PDF using OCR"""
        try:
            # Convert PDF pages to images
            images = pdf2image.convert_from_path(
                pdf_path,
                dpi=200
            )
            
            extracted_text = ""
            
            for i, image in enumerate(images):
                # Save image temporarily
                img_path = os.path.join(self.temp_dir, f"temp_ocr_page_{i}_{int(time.time())}.png")
                image.save(img_path)
                
                try:
                    # Auto-detect language if needed
                    if language == "auto_detect":
                        tesseract_lang = self.detect_language(img_path)
                    else:
                        tesseract_lang = self.get_tesseract_language(language)
                    
                    # Perform OCR
                    page_text = pytesseract.image_to_string(
                        img_path,
                        lang=tesseract_lang,
                        config='--oem 3 --psm 6'
                    )
                    
                    if page_text.strip():
                        extracted_text += f"--- Page {i+1} ---\n{page_text}\n\n"
                
                except Exception as e:
                    extracted_text += f"--- Page {i+1} ---\nError extracting text: {str(e)}\n\n"
                
                finally:
                    # Clean up temporary image
                    if os.path.exists(img_path):
                        os.remove(img_path)
            
            # Save extracted text
            if output_format == "plain text":
                output_path = os.path.join(self.temp_dir, f"ocr_extracted_text_{int(time.time())}.txt")
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(extracted_text)
            else:  # Word document
                doc = Document()
                doc.add_heading('OCR Extracted Text', 0)
                
                for paragraph in extracted_text.split('\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                
                output_path = os.path.join(self.temp_dir, f"ocr_extracted_text_{int(time.time())}.docx")
                doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error extracting text with OCR: {str(e)}")
    
    def enhance_image_for_ocr(self, image_path: str) -> str:
        """Enhance image quality for better OCR results"""
        try:
            from PIL import Image, ImageEnhance, ImageFilter
            
            # Open and process image
            image = Image.open(image_path)
            
            # Convert to grayscale for better OCR
            if image.mode != 'L':
                image = image.convert('L')
            
            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)
            
            # Enhance sharpness
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(2.0)
            
            # Apply slight blur to reduce noise
            image = image.filter(ImageFilter.MedianFilter(size=3))
            
            # Save enhanced image
            enhanced_path = image_path.replace('.png', '_enhanced.png')
            image.save(enhanced_path, 'PNG', optimize=True)
            
            return enhanced_path
            
        except Exception as e:
            # Return original path if enhancement fails
            return image_path

import time
