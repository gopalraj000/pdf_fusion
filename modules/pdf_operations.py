import os
import tempfile
import zipfile
from typing import List, Optional, Tuple
import fitz  # PyMuPDF
import PyPDF2
from PyPDF2 import PdfWriter, PdfReader
# import pdfplumber  # Removed due to compatibility issues
from io import BytesIO
import re

class PDFOperations:
    """Handles all PDF-related operations including merge, split, compress, etc."""
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def get_page_count(self, pdf_path: str) -> int:
        """Get the number of pages in a PDF"""
        try:
            with fitz.open(pdf_path) as doc:
                return len(doc)
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def merge_pdfs(self, pdf_paths: List[str], add_blank_pages: bool = False, add_bookmarks: bool = True) -> Optional[str]:
        """Merge multiple PDFs into one"""
        try:
            output_path = os.path.join(self.temp_dir, f"merged_{int(time.time())}.pdf")
            writer = PdfWriter()
            
            for i, pdf_path in enumerate(pdf_paths):
                reader = PdfReader(pdf_path)
                num_pages = len(reader.pages)
                
                # Add bookmark for this file
                if add_bookmarks:
                    bookmark_name = os.path.basename(pdf_path).replace('.pdf', '')
                    writer.add_outline_item(bookmark_name, len(writer.pages))
                
                # Add all pages from this PDF
                for page in reader.pages:
                    writer.add_page(page)
                
                # Add blank page if this PDF has odd number of pages and option is enabled
                if add_blank_pages and num_pages % 2 == 1 and i < len(pdf_paths) - 1:
                    blank_page = writer.add_blank_page()
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error merging PDFs: {str(e)}")
    
    def split_pdf_by_ranges(self, pdf_path: str, ranges_str: str) -> List[str]:
        """Split PDF by specified page ranges"""
        try:
            reader = PdfReader(pdf_path)
            total_pages = len(reader.pages)
            output_paths = []
            
            # Parse ranges (e.g., "1-5, 8-12, 15-20")
            ranges = []
            for range_part in ranges_str.split(','):
                range_part = range_part.strip()
                if '-' in range_part:
                    start, end = map(int, range_part.split('-'))
                    ranges.append((start - 1, end - 1))  # Convert to 0-based indexing
                else:
                    page_num = int(range_part) - 1
                    ranges.append((page_num, page_num))
            
            for i, (start, end) in enumerate(ranges):
                if start < 0 or end >= total_pages or start > end:
                    raise ValueError(f"Invalid page range: {start+1}-{end+1}")
                
                writer = PdfWriter()
                for page_num in range(start, end + 1):
                    writer.add_page(reader.pages[page_num])
                
                output_path = os.path.join(self.temp_dir, f"split_part_{i+1}_{int(time.time())}.pdf")
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                
                output_paths.append(output_path)
            
            return output_paths
            
        except Exception as e:
            raise Exception(f"Error splitting PDF: {str(e)}")
    
    def extract_specific_pages(self, pdf_path: str, pages_str: str) -> Optional[str]:
        """Extract specific pages from PDF"""
        try:
            reader = PdfReader(pdf_path)
            total_pages = len(reader.pages)
            writer = PdfWriter()
            
            # Parse page numbers (e.g., "1, 3, 5, 8")
            page_numbers = []
            for page_str in pages_str.split(','):
                page_num = int(page_str.strip()) - 1  # Convert to 0-based indexing
                if 0 <= page_num < total_pages:
                    page_numbers.append(page_num)
                else:
                    raise ValueError(f"Invalid page number: {page_num + 1}")
            
            # Sort page numbers to maintain order
            page_numbers.sort()
            
            for page_num in page_numbers:
                writer.add_page(reader.pages[page_num])
            
            output_path = os.path.join(self.temp_dir, f"extracted_pages_{int(time.time())}.pdf")
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error extracting pages: {str(e)}")
    
    def split_pdf_equal_parts(self, pdf_path: str, num_parts: int) -> List[str]:
        """Split PDF into equal parts"""
        try:
            reader = PdfReader(pdf_path)
            total_pages = len(reader.pages)
            pages_per_part = total_pages // num_parts
            output_paths = []
            
            for i in range(num_parts):
                writer = PdfWriter()
                start_page = i * pages_per_part
                
                # Last part gets any remaining pages
                if i == num_parts - 1:
                    end_page = total_pages
                else:
                    end_page = (i + 1) * pages_per_part
                
                for page_num in range(start_page, end_page):
                    writer.add_page(reader.pages[page_num])
                
                output_path = os.path.join(self.temp_dir, f"part_{i+1}_{int(time.time())}.pdf")
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                
                output_paths.append(output_path)
            
            return output_paths
            
        except Exception as e:
            raise Exception(f"Error splitting PDF into parts: {str(e)}")
    
    def compress_pdf(self, pdf_path: str, compression_level: str = "medium") -> Optional[str]:
        """Compress PDF to reduce file size"""
        try:
            doc = fitz.open(pdf_path)
            
            # Compression settings based on level
            if compression_level == "low":
                deflate_level = 1
                image_quality = 85
            elif compression_level == "medium":
                deflate_level = 6
                image_quality = 70
            elif compression_level == "high":
                deflate_level = 9
                image_quality = 50
            else:  # maximum
                deflate_level = 9
                image_quality = 30
            
            # Process each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Get images on the page and compress them
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # Compress image
                    if image_ext in ["png", "jpg", "jpeg"]:
                        from PIL import Image
                        import io
                        
                        pil_image = Image.open(io.BytesIO(image_bytes))
                        if pil_image.mode == "RGBA":
                            pil_image = pil_image.convert("RGB")
                        
                        # Compress and save
                        compressed_bytes = io.BytesIO()
                        pil_image.save(compressed_bytes, format="JPEG", quality=image_quality, optimize=True)
                        compressed_image_bytes = compressed_bytes.getvalue()
                        
                        # Replace image in PDF
                        doc.update_stream(xref, compressed_image_bytes)
            
            output_path = os.path.join(self.temp_dir, f"compressed_{int(time.time())}.pdf")
            doc.save(output_path, deflate=True)
            doc.close()
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error compressing PDF: {str(e)}")
    
    def reorder_pages(self, pdf_path: str, new_order_str: str) -> Optional[str]:
        """Reorder pages in PDF based on specified order"""
        try:
            reader = PdfReader(pdf_path)
            total_pages = len(reader.pages)
            writer = PdfWriter()
            
            # Parse new order (e.g., "3,1,2,4,5")
            new_order = []
            for page_str in new_order_str.split(','):
                page_num = int(page_str.strip()) - 1  # Convert to 0-based indexing
                if 0 <= page_num < total_pages:
                    new_order.append(page_num)
                else:
                    raise ValueError(f"Invalid page number: {page_num + 1}")
            
            # Add pages in new order
            for page_num in new_order:
                writer.add_page(reader.pages[page_num])
            
            output_path = os.path.join(self.temp_dir, f"reordered_{int(time.time())}.pdf")
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error reordering pages: {str(e)}")
    
    def add_password_protection(self, pdf_path: str, user_password: str, owner_password: str, permissions: List[str]) -> Optional[str]:
        """Add password protection to PDF"""
        try:
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            
            # Add all pages
            for page in reader.pages:
                writer.add_page(page)
            
            # Set permissions
            allow_printing = "Print" in permissions
            allow_copying = "Copy Text" in permissions
            allow_modifying = "Modify" in permissions
            allow_annotations = "Add Annotations" in permissions
            
            # Encrypt PDF
            writer.encrypt(
                user_password=user_password,
                owner_password=owner_password,
                use_128bit=True
            )
            
            output_path = os.path.join(self.temp_dir, f"protected_{int(time.time())}.pdf")
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error adding password protection: {str(e)}")
    
    def remove_password_protection(self, pdf_path: str, password: str) -> Optional[str]:
        """Remove password protection from PDF"""
        try:
            reader = PdfReader(pdf_path)
            
            # Decrypt PDF
            if reader.is_encrypted:
                if not reader.decrypt(password):
                    raise ValueError("Incorrect password")
            
            writer = PdfWriter()
            
            # Add all pages
            for page in reader.pages:
                writer.add_page(page)
            
            output_path = os.path.join(self.temp_dir, f"unlocked_{int(time.time())}.pdf")
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error removing password protection: {str(e)}")
    
    def extract_images(self, pdf_path: str, format: str = "png", min_size: int = 100) -> List[str]:
        """Extract images from PDF"""
        try:
            doc = fitz.open(pdf_path)
            image_paths = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # Check image size
                    from PIL import Image
                    import io
                    
                    pil_image = Image.open(io.BytesIO(image_bytes))
                    if pil_image.width < min_size or pil_image.height < min_size:
                        continue
                    
                    # Save image
                    output_path = os.path.join(
                        self.temp_dir, 
                        f"extracted_image_p{page_num+1}_{img_index+1}_{int(time.time())}.{format}"
                    )
                    
                    if format.lower() == "png":
                        pil_image.save(output_path, "PNG")
                    else:  # JPEG
                        if pil_image.mode == "RGBA":
                            pil_image = pil_image.convert("RGB")
                        pil_image.save(output_path, "JPEG", quality=95)
                    
                    image_paths.append(output_path)
            
            doc.close()
            return image_paths
            
        except Exception as e:
            raise Exception(f"Error extracting images: {str(e)}")
    
    def extract_text(self, pdf_path: str, output_format: str = "plain text") -> Optional[str]:
        """Extract text from PDF"""
        try:
            extracted_text = ""
            
            # Use PyMuPDF for text extraction
            doc = fitz.open(pdf_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                if page_text:
                    extracted_text += f"--- Page {page_num + 1} ---\n{page_text}\n\n"
            doc.close()
            
            if output_format == "plain text":
                output_path = os.path.join(self.temp_dir, f"extracted_text_{int(time.time())}.txt")
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(extracted_text)
            else:  # Word document
                from docx import Document
                
                doc = Document()
                doc.add_heading('Extracted Text', 0)
                for paragraph in extracted_text.split('\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                
                output_path = os.path.join(self.temp_dir, f"extracted_text_{int(time.time())}.docx")
                doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error extracting text: {str(e)}")
    
    def create_images_zip(self, image_paths: List[str], original_filename: str) -> str:
        """Create a ZIP file containing all extracted images"""
        try:
            zip_path = os.path.join(self.temp_dir, f"{original_filename}_images_{int(time.time())}.zip")
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for i, image_path in enumerate(image_paths):
                    filename = f"image_{i+1:03d}.{image_path.split('.')[-1]}"
                    zipf.write(image_path, filename)
            
            return zip_path
            
        except Exception as e:
            raise Exception(f"Error creating images ZIP: {str(e)}")

import time
